"""
Document parsing service for text extraction from various file formats.
Supports: Plain text, PDF, Word (.docx)
"""

import io
import logging
import re
from pathlib import Path
from typing import BinaryIO

import chardet
from docx import Document as DocxDocument
from pypdf import PdfReader

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for extracting text from various document formats."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx", ".doc", ".md"}

    def parse(self, file: BinaryIO, filename: str) -> str:
        """
        Parse a file and extract text content.

        Args:
            file: File-like object containing the document
            filename: Original filename (used to determine format)

        Returns:
            Extracted text content
        """
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file)
        elif ext in (".txt", ".md"):
            return self._parse_text(file)
        else:
            # Try to parse as text
            return self._parse_text(file)

    def parse_text(self, text: str) -> str:
        """Parse raw text content."""
        return self._clean_text(text)

    def _parse_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file)
            text_parts = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return self._clean_text("\n\n".join(text_parts))
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    def _parse_docx(self, file: BinaryIO) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(file)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return self._clean_text("\n\n".join(text_parts))
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _parse_text(self, file: BinaryIO) -> str:
        """Extract text from plain text file."""
        try:
            content = file.read()

            # Detect encoding
            if isinstance(content, bytes):
                detected = chardet.detect(content)
                encoding = detected.get("encoding", "utf-8") or "utf-8"
                text = content.decode(encoding, errors="replace")
            else:
                text = content

            return self._clean_text(text)
        except Exception as e:
            logger.error(f"Failed to parse text file: {e}")
            raise ValueError(f"Failed to parse text file: {e}")

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        # Remove control characters except newlines
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        return text.strip()


class TextChunker:
    """Split text into chunks for processing."""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: list[str] | None = None,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", "。", ".", " ", ""]

    def split(self, text: str) -> list[str]:
        """
        Split text into chunks.

        Args:
            text: Text to split

        Returns:
            List of text chunks
        """
        if not text:
            return []

        chunks = self._recursive_split(text, self.separators)

        # Merge small chunks and handle overlap
        final_chunks = []
        current_chunk = ""

        for chunk in chunks:
            if len(current_chunk) + len(chunk) <= self.chunk_size:
                current_chunk += chunk
            else:
                if current_chunk:
                    final_chunks.append(current_chunk.strip())
                    # Keep overlap from end of current chunk
                    if self.chunk_overlap > 0:
                        overlap_text = current_chunk[-self.chunk_overlap:]
                        current_chunk = overlap_text + chunk
                    else:
                        current_chunk = chunk
                else:
                    current_chunk = chunk

        if current_chunk.strip():
            final_chunks.append(current_chunk.strip())

        return final_chunks

    def _recursive_split(self, text: str, separators: list[str]) -> list[str]:
        """Recursively split text using separators."""
        if not separators:
            return [text]

        sep = separators[0]
        remaining_seps = separators[1:]

        if sep == "":
            # Character-level split
            return [text[i : i + self.chunk_size] for i in range(0, len(text), self.chunk_size)]

        parts = text.split(sep)
        result = []

        for part in parts:
            if len(part) <= self.chunk_size:
                result.append(part + sep if sep else part)
            else:
                # Recursively split large parts
                sub_parts = self._recursive_split(part, remaining_seps)
                result.extend(sub_parts)

        return result


def get_document_parser() -> DocumentParser:
    """Get document parser instance."""
    return DocumentParser()


def get_text_chunker(
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> TextChunker:
    """Get text chunker instance."""
    return TextChunker(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
